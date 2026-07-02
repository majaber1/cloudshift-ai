from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc
from database import get_db, Assessment, AssessmentDocument, User
from routers.auth import get_current_user
from pydantic import BaseModel
from typing import Optional, List
import uuid
import os
import aiofiles
import PyPDF2
import docx
import io

router = APIRouter()
UPLOAD_DIR = "/app/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class AssessmentCreate(BaseModel):
    name: str
    organization: str
    description: Optional[str] = None

@router.post("/")
async def create_assessment(
    assessment_data: AssessmentCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    assessment = Assessment(
        name=assessment_data.name,
        organization=assessment_data.organization,
        description=assessment_data.description,
        user_id=current_user.id
    )
    db.add(assessment)
    await db.commit()
    await db.refresh(assessment)
    return {"id": assessment.id, "name": assessment.name, "organization": assessment.organization, "status": assessment.status}

@router.get("/")
async def list_assessments(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    result = await db.execute(
        select(Assessment).where(Assessment.user_id == current_user.id).order_by(desc(Assessment.created_at))
    )
    assessments = result.scalars().all()
    return [{"id": a.id, "name": a.name, "organization": a.organization, "status": a.status, "created_at": str(a.created_at)} for a in assessments]

@router.get("/{assessment_id}")
async def get_assessment(
    assessment_id: int,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    result = await db.execute(select(Assessment).where(Assessment.id == assessment_id, Assessment.user_id == current_user.id))
    assessment = result.scalar_one_or_none()
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    docs_result = await db.execute(select(AssessmentDocument).where(AssessmentDocument.assessment_id == assessment_id))
    documents = docs_result.scalars().all()
    
    return {
        "id": assessment.id,
        "name": assessment.name,
        "organization": assessment.organization,
        "description": assessment.description,
        "status": assessment.status,
        "created_at": str(assessment.created_at),
        "documents": [{"id": d.id, "filename": d.original_filename, "doc_type": d.doc_type} for d in documents]
    }

@router.post("/{assessment_id}/upload")
async def upload_document(
    assessment_id: int,
    doc_type: str = Form(...),
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    result = await db.execute(select(Assessment).where(Assessment.id == assessment_id, Assessment.user_id == current_user.id))
    assessment = result.scalar_one_or_none()
    if not assessment:
        raise HTTPException(status_code=404, detail="Assessment not found")
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    content = await file.read()
    async with aiofiles.open(file_path, 'wb') as f:
        await f.write(content)
    
    extracted_text = ""
    try:
        if file_ext == '.pdf':
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            extracted_text = " ".join([page.extract_text() or "" for page in reader.pages])
        elif file_ext in ['.docx', '.doc']:
            doc = docx.Document(io.BytesIO(content))
            extracted_text = " ".join([para.text for para in doc.paragraphs])
        else:
            extracted_text = content.decode('utf-8', errors='ignore')
    except Exception as e:
        extracted_text = f"Could not extract text: {str(e)}"
    
    document = AssessmentDocument(
        assessment_id=assessment_id,
        doc_type=doc_type,
        original_filename=file.filename,
        stored_filename=unique_filename,
        extracted_text=extracted_text[:50000]
    )
    db.add(document)
    await db.commit()
    await db.refresh(document)
    
    return {"id": document.id, "filename": file.filename, "doc_type": doc_type, "text_length": len(extracted_text)}
